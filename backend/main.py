from fastapi import FastAPI
from database import engine
import models
from auth import router as auth_router
from fastapi import Depends
from auth import get_current_user
import models
from fastapi import File, UploadFile
from sqlalchemy.orm import Session
from auth import get_current_user
from database import SessionLocal
import models
from rag import simple_keyword_match, generate_answer_from_context
from fastapi.responses import FileResponse
from docx import Document
import os

models.Base.metadata.create_all(bind=engine)

app = FastAPI()

app.include_router(auth_router)

@app.get("/")
def root():
    return {"message": "Auth working"}



@app.get("/protected")
def protected_route(current_user: models.User = Depends(get_current_user)):
    return {"message": f"Hello {current_user.email}"}

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@app.post("/upload-questionnaire")
async def upload_questionnaire(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    content = await file.read()
    text_content = content.decode("utf-8")

    print("RAW TEXT CONTENT:")
    print(text_content)

    lines = text_content.split("\n")
    print("SPLIT LINES:", lines)

    questionnaire = models.Questionnaire(
        filename=file.filename,
        content=text_content,
        owner_id=current_user.id
    )

    db.add(questionnaire)
    db.commit()
    db.refresh(questionnaire)

    for line in lines:
        cleaned = line.strip()
        if cleaned:
            print("ADDING QUESTION:", cleaned)
            question = models.Question(
                text=cleaned,
                questionnaire_id=questionnaire.id
            )
            db.add(question)

    db.commit()

    return {"message": "Questionnaire uploaded and parsed successfully"}

@app.post("/upload-reference")
async def upload_reference(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    content = await file.read()
    text_content = content.decode("utf-8")

    reference = models.ReferenceDocument(
        filename=file.filename,
        content=text_content,
        owner_id=current_user.id
    )

    db.add(reference)
    db.commit()
    db.refresh(reference)

    return {"message": "Reference document uploaded successfully"}

@app.post("/generate-answers")
def generate_answers(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.owner_id == current_user.id
    ).order_by(models.Questionnaire.id.desc()).first()

    if not questionnaire:
        return {"error": "No questionnaire found"}

    references = db.query(models.ReferenceDocument).filter(
        models.ReferenceDocument.owner_id == current_user.id
    ).all()

    if not references:
        return {"error": "No reference documents found"}

    documents = [
        {"text": ref.content, "filename": ref.filename}
        for ref in references
    ]

    questions = db.query(models.Question).filter(
        models.Question.questionnaire_id == questionnaire.id
    ).all()

    results = []
    answered_count = 0
    not_found_count = 0

    for q in questions:
        matched_docs = simple_keyword_match(q.text, documents)
        answer_text, citation, confidence = generate_answer_from_context(q.text, matched_docs)

        if answer_text == "Not found in references.":
            not_found_count += 1
        else:
            answered_count += 1

        answer = models.Answer(
            text=answer_text,
            citation=citation if citation else "N/A",
            question_id=q.id
        )

        db.add(answer)

        results.append({
            "question": q.text,
            "answer": answer_text,
            "citation": citation if citation else "N/A",
            "confidence": confidence
        })

    db.commit()

    summary = {
        "total_questions": len(questions),
        "answered": answered_count,
        "not_found": not_found_count
    }

    return {
        "summary": summary,
        "results": results
    }

@app.get("/debug/questionnaires")
def debug_questionnaires(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    data = db.query(models.Questionnaire).filter(
        models.Questionnaire.owner_id == current_user.id
    ).all()

    return [{"id": q.id, "filename": q.filename} for q in data]


@app.get("/debug/questions")
def debug_questions(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    data = db.query(models.Question).all()

    return [{"id": q.id, "text": q.text} for q in data]


@app.get("/debug/references")
def debug_references(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    data = db.query(models.ReferenceDocument).filter(
        models.ReferenceDocument.owner_id == current_user.id
    ).all()

    return [{"id": r.id, "filename": r.filename} for r in data]


@app.post("/export")
def export_document(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.owner_id == current_user.id
    ).order_by(models.Questionnaire.id.desc()).first()

    if not questionnaire:
        return {"error": "No questionnaire found"}

    questions = db.query(models.Question).filter(
        models.Question.questionnaire_id == questionnaire.id
    ).all()

    document = Document()

    for q in questions:
        document.add_paragraph(q.text)

        answer = db.query(models.Answer).filter(
            models.Answer.question_id == q.id
        ).order_by(models.Answer.id.desc()).first()

        if answer:
            document.add_paragraph(f"Answer: {answer.text}")
            document.add_paragraph(f"Citation: {answer.citation}")
        else:
            document.add_paragraph("Answer: Not generated.")
        
        document.add_paragraph("")

    file_path = "generated_answers.docx"
    document.save(file_path)

    return FileResponse(
        path=file_path,
        filename="structured_answers.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )